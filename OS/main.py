
#   Copyright (c) 2026 Teodor Nenkov

#   Licensed under the PolyForm Noncommercial License 1.0.0.
#   Commercial use requires a separate license.

#   See LICENSE for details.

#   Europe, Bulgaria

import asyncio
import base64
import json
import os
import re
import shutil
import urllib.parse
import subprocess
import tempfile
import openpyxl
import uuid
import docx
import zipfile
import subprocess
import sys
import glob
import io
import stat
import gc
import threading
import logging
import requests
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime
from io import BytesIO
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile, WebSocket, Query, BackgroundTasks, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from playwright.async_api import async_playwright
from pydantic import BaseModel
from logic import OlachraEngine



app = FastAPI(title="OLACHRA")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

os.environ['POSTHOG_DISABLED'] = '1'
os.environ['ANONYMIZED_TELEMETRY'] = 'False'
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
KNOWLEDGE_DIR = os.path.join(BASE_DIR, "documents")
DOC_EXTERNAL_DIR = os.path.join(KNOWLEDGE_DIR, "external")
DOC_INTERNAL_DIR = os.path.join(KNOWLEDGE_DIR, "internal")
DOC_META_FILE = os.path.join(KNOWLEDGE_DIR, "documents.json")
SESSIONS_DIR = os.path.join(BASE_DIR, "sessions_json")
user_data_dir = os.path.join(BASE_DIR, "browser_user_data")
ASTRAL_DIR = "astral_projection"
CLIPBOARD_DIR = "clipboard"
CLIPBOARD_FILE = os.path.join(CLIPBOARD_DIR, "sessions_clipboard.json")
ASTRAL_FILE = os.path.join(ASTRAL_DIR, "astral_data.json")
PLANCHETTE_DIR = os.path.join(DOC_INTERNAL_DIR, "planchette_mode")
SCRIPTORIA_DIR = os.path.join(BASE_DIR, "scriptoria")
SCRIPTORIA_NOTES = os.path.join(SCRIPTORIA_DIR, "notes.json")
SCRIPTORIA_DATASET = os.path.join(SCRIPTORIA_DIR, "dataset.jsonl")
CONVERTERS_DIR = os.path.join(SCRIPTORIA_DIR, "converters")
SCRIPTORIA_DOCX = os.path.join(CONVERTERS_DIR, "docx_files")
JSONL_FILES_DIR = os.path.join(CONVERTERS_DIR, "jsonl_files")
JSONL_MERGED_DIR = os.path.join(CONVERTERS_DIR, "jsonl_merged_files")
JSONL_L_DIR = os.path.join(JSONL_FILES_DIR, "L")
JSONL_U_DIR = os.path.join(JSONL_FILES_DIR, "U")
JSONL_W_DIR = os.path.join(JSONL_FILES_DIR, "W")

TEMP_PASSIVE_SESSIONS = set()

CONFIG_DIRS = {
    "personalities": os.path.join(BASE_DIR, "personalities"),
    "emotionalities": os.path.join(BASE_DIR, "emotionalities"),
    "values": os.path.join(BASE_DIR, "value_systems"),
    "serper_keys": os.path.join(BASE_DIR, "serper_keys"),
    "cognitive_layers": os.path.join(BASE_DIR, "cognitive_layers")
}

ALL_DIRS = [
    KNOWLEDGE_DIR, SESSIONS_DIR, DOC_EXTERNAL_DIR, DOC_INTERNAL_DIR, 
    ASTRAL_DIR, CLIPBOARD_DIR, PLANCHETTE_DIR, SCRIPTORIA_DIR, 
    CONVERTERS_DIR, SCRIPTORIA_DOCX, JSONL_FILES_DIR, JSONL_MERGED_DIR,
    JSONL_L_DIR, JSONL_U_DIR, JSONL_W_DIR
] + list(CONFIG_DIRS.values())

for d in ALL_DIRS:
    os.makedirs(d, exist_ok=True)


ALLOWED_ROOT_FOLDERS = [
    "astral_projection", 
    "clipboard", 
    "cognitive_layers", 
    "documents",
    "emotionalities",
    "personalities", 
    "value_systems", 
    "sessions_json",
    "serper_keys", 
    "scriptoria"
]


if not os.path.exists(CLIPBOARD_FILE) or os.path.getsize(CLIPBOARD_FILE) == 0:
    with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
        json.dump([], f)

if not os.path.exists(ASTRAL_FILE) or os.path.getsize(ASTRAL_FILE) == 0:
    with open(ASTRAL_FILE, "w", encoding="utf-8") as f:
        json.dump({"cells": []}, f)


engine = OlachraEngine()

app.mount("/static", StaticFiles(directory="static"), name="static")


# Returns a successful 204 No Content response for the favicon icon request
@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    return Response(status_code=204)

# Returns a successful 204 No Content response for static library map file requests
@app.get("/static/libraries/{file_path}.map", include_in_schema=False)
async def ignore_maps():
    return Response(status_code=204)

# Returns a successful 204 No Content response for Chrome DevTools JSON requests
@app.get("/.well-known/appspecific/com.chrome.devtools.json", include_in_schema=False)
async def ignore_chrome_json():
    return Response(status_code=204)

# Clears accumulated gigabytes of browser cache upon restart while preserving cookies
def clear_chromium_heavy_cache():
    """Clears gigabytes of accumulated browser cache on restart, but keeps cookies."""
    cache_paths = [
        os.path.join(user_data_dir, "Default", "Cache"),
        os.path.join(user_data_dir, "Default", "Code Cache"),
        os.path.join(user_data_dir, "Default", "GPUCache")
    ]
    for path in cache_paths:
        if os.path.exists(path):
            try:
                shutil.rmtree(path)
                print(f"Cleaned browser cache: {os.path.basename(path)}")
            except Exception:
                pass

clear_chromium_heavy_cache()

if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

from asyncio.log import logger
logger.addFilter(lambda record: "ProactorBasePipeTransport" not in record.getMessage())

def normalize_path(path_str: str) -> str:
    """Universal normalizer: converts Windows paths (\\) to Linux paths (/)"""
    if not path_str: return ""
    return path_str.replace("\\", "/")

# Triggered on application startup to clean up things
@app.on_event("startup")
async def startup_event():
    print("Check for pending temporary sessions...")
    if os.path.exists(SESSIONS_DIR):
        for f in os.listdir(SESSIONS_DIR):
            if f.startswith("temp_passive_") and f.endswith(".json"):
                try:
                    os.remove(os.path.join(SESSIONS_DIR, f))
                except: pass

    for layer in ["citta", "vritti"]:
        try:
            layer_path = os.path.join(engine.layers_dir, engine.LAYER_MAP[layer], f"{layer}.json")
            with open(layer_path, 'w', encoding='utf-8') as f:
                json.dump([], f)
            print(f"Layer {layer.upper()} is completely cleaned at start.")
        except Exception:
            pass

# Triggered on application shutdown to clean up active temporary passive sessions
@app.on_event("shutdown")
def cleanup_temp_sessions():
    print("Clean up active temporary passive sessions...")
    for file_path in TEMP_PASSIVE_SESSIONS:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except: pass

# Checks and returns the last modified timestamp of a given session file
@app.get("/api/sessions/check-modified")
async def check_modified(file: str):
    if not file or "/" in file or "\\" in file:
        return {"last_modified": 0}
        
    file_path = os.path.join("sessions_json", file) 
    
    try:
        mtime = os.path.getmtime(file_path)
        return {"last_modified": mtime}
    except Exception:
        return {"last_modified": 0}

class InlineChatRequest(BaseModel):
    type: str
    target_text: str
    instruction: str
    model: str
    mode: int
    session_file: Optional[str] = None
    text_before: Optional[str] = "" 
    text_after: Optional[str] = ""
    system_layers: List[str] = [] 

class ChatRequest(BaseModel):
    message: str
    mode: int                   
    session_file: Optional[str] = None
    combined_sessions: List[str] = [] 
    model: str
    personality: str
    emotion: str
    values: List[str]
    selected_files: List[str]  
    web_search: bool 
    insight: bool = False     
    temperature: float = 0.7  
    serper_key: Optional[str] = None
    browser_context: str = ""
    system_layers: List[str] = [] 
    
class RenameRequest(BaseModel):
    filename: str
    new_name: str

class ConfigItem(BaseModel):
    name: str
    system: str

class DeleteMessageRequest(BaseModel):
    filename: str
    msg_index: int

# Renumbers sessions sequentially based on their creation time to maintain consistent naming logic
def reindex_sessions():
    """
    Renumbers sessions sequentially based on creation time.
    Ensures Session #1, #2, etc. logic remains consistent.
    """
    if not os.path.exists(SESSIONS_DIR): return
    
    sessions = []
    files = [f for f in os.listdir(SESSIONS_DIR) if f.endswith(".json") and not f.startswith("temp_passive_")]
    
    for f in files:
        path = os.path.join(SESSIONS_DIR, f)
        try:
            with open(path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                created_at = data.get("created_at", str(datetime.now()))
                sessions.append({"data": data, "created_at": created_at})
        except: pass
    
    sessions.sort(key=lambda x: x["created_at"])
    
    for f in files:
        try:
            os.remove(os.path.join(SESSIONS_DIR, f))
        except: pass

    for index, session in enumerate(sessions):
        new_id = index + 1
        new_filename = f"session_{new_id}.json"
        
        session_data = session["data"]
        session_data["id"] = new_id
        
        if session_data["name"].startswith("Session #") or session_data["name"] == "Session":
             session_data["name"] = f"Session #{new_id}"
        
        new_path = os.path.join(SESSIONS_DIR, new_filename)
        
        with open(new_path, 'w', encoding='utf-8') as f:
            json.dump(session_data, f, indent=4, ensure_ascii=False)

# Reads and returns the content from the specified configuration JSON files
def load_sys_content(category, name):
    """Reads content from config JSON files."""
    if not name: return ""
    folder = CONFIG_DIRS.get(category)
    if not folder: return ""
    
    filename = name if name.endswith(".json") else f"{name}.json"
    path = os.path.join(folder, filename)
    
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return data.get('system', '')
    except:
        return ""

# Synchronizes the physical files in the knowledge directory with the documents metadata JSON
def sync_documents_meta():
    """Syncs the actual files in documents/ with documents.json"""
    DOC_META_FILE = os.path.join(KNOWLEDGE_DIR, "documents.json")
    
    if not os.path.exists(DOC_META_FILE):
        meta = {"external": [], "internal": {}}
    else:
        try:
            with open(DOC_META_FILE, 'r', encoding='utf-8') as f:
                meta = json.load(f)
        except Exception:
            meta = {"external": [], "internal": {}}

    existing_ext = {item['name']: item for item in meta.get('external', [])}
    actual_ext = []
    
    if os.path.exists(KNOWLEDGE_DIR):
        for fname in os.listdir(KNOWLEDGE_DIR):
            fpath = os.path.join(KNOWLEDGE_DIR, fname)
            if os.path.isfile(fpath) and fname != "documents.json":
                if fname in existing_ext:
                    actual_ext.append(existing_ext[fname])
                else:
                    actual_ext.append({
                        "name": fname,
                        "path": f"documents/{fname}",
                        "size": os.path.getsize(fpath),
                        "added_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "type": "file"
                    })
    meta['external'] = actual_ext

    with open(DOC_META_FILE, 'w', encoding='utf-8') as f:
        json.dump(meta, f, indent=4, ensure_ascii=False)
        
    return meta

# Returns a specially grouped hierarchical structure of documents for the left panel interface
@app.get("/api/fs/documents")
def get_grouped_documents():
    """Returns the special grouped structure for the left panel"""
    meta = sync_documents_meta()
    meta["is_grouped"] = True
    return meta

# Endpoint for securely loading and returning the contents of plain text files
@app.get("/api/fs/read")
def read_fs_file(path: str):
    """Endpoint for loading only plain text files"""
    path = normalize_path(path)
    if '..' in path: raise HTTPException(status_code=403, detail="Invalid path")
    full_path = os.path.join(BASE_DIR, path)
    if not os.path.exists(full_path): raise HTTPException(status_code=404, detail="File not found")
    
    encodings_to_try = ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']
    for enc in encodings_to_try:
        try:
            with open(full_path, 'r', encoding=enc) as f:
                return {"content": f.read()}
        except UnicodeDecodeError:
            continue
    
    raise HTTPException(status_code=500, detail="This format is not supported for direct viewing.")

class SaveDirectRequest(BaseModel):
    path: str
    content: str

# Directly saves content to a specific file path after validating against directory traversal
@app.post("/api/fs/save_direct")
def save_direct_doc(req: SaveDirectRequest):
    req.path = normalize_path(req.path)
    if '..' in req.path: raise HTTPException(status_code=403, detail="Invalid path")
    full_path = os.path.join(BASE_DIR, req.path)
    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    with open(full_path, 'w', encoding='utf-8') as f:
        f.write(req.content)
        
    return {"status": "success"}

class SaveInternalRequest(BaseModel):
    original_name: str
    content: str
    session_name: str 

# Saves an internal document by synchronizing document metadata and writing the content
@app.post("/api/fs/save_internal")
def save_internal_doc(req: SaveInternalRequest):
    meta = sync_documents_meta()
    
    session_folder_name = req.session_name.replace(" ", "_").replace("#", "").lower()
    DOC_INTERNAL_DIR = os.path.join(KNOWLEDGE_DIR, "internal")
    session_dir = os.path.join(DOC_INTERNAL_DIR, session_folder_name)
    os.makedirs(session_dir, exist_ok=True)
    
    base, ext = os.path.splitext(req.original_name)
    new_filename = f"{base}_edited{ext}" if not base.endswith("_edited") else req.original_name
    file_path = os.path.join(session_dir, new_filename)
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(req.content)
        
    if "internal" not in meta: meta["internal"] = {}
    if req.session_name not in meta["internal"]: meta["internal"][req.session_name] = []
        
    file_record = next((f for f in meta["internal"][req.session_name] if f["name"] == new_filename), None)
    
    if not file_record:
        meta["internal"][req.session_name].append({
            "name": new_filename,
            "path": f"documents/internal/{session_folder_name}/{new_filename}",
            "size": os.path.getsize(file_path),
            "edited_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "type": "file"
        })
    else:
        file_record["size"] = os.path.getsize(file_path)
        file_record["edited_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        
    DOC_META_FILE = os.path.join(KNOWLEDGE_DIR, "documents.json")
    with open(DOC_META_FILE, 'w', encoding='utf-8') as f:
        json.dump(meta, f, indent=4, ensure_ascii=False)
        
    return {"status": "success", "path": f"documents/internal/{session_folder_name}/{new_filename}"}

# Retrieves configuration items from a specified category folder
@app.get("/api/config/{category}")
def get_configs(category: str):
    folder = CONFIG_DIRS.get(category)
    if not folder: return {"items": []}
    
    items = []
    for root, _, files in os.walk(folder):
        for f in files:
            if f.endswith(".json"):
                if f != "hall_of_maat.json":
                    items.append(f.replace(".json", ""))
                    
    return {"items": sorted(items)}

# Creates a new configuration item within a specified category folder
@app.post("/api/config/{category}")
def create_config(category: str, item: ConfigItem):
    folder = CONFIG_DIRS.get(category)
    if not folder: raise HTTPException(status_code=400, detail="Invalid category")
    
    filename = f"{item.name.strip().replace(' ', '_')}.json"
    path = os.path.join(folder, filename)
    data = {"name": item.name, "system": item.system}
    
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        return {"status": "success", "filename": filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Deletes a specific configuration item from a designated category folder
@app.delete("/api/config/{category}/{item_name}")
def delete_config(category: str, item_name: str):
    folder = CONFIG_DIRS.get(category)
    if not folder: raise HTTPException(status_code=400, detail="Invalid category")
    
    safe_name = os.path.basename(item_name)
    filename = f"{safe_name}.json"
    path = os.path.join(folder, filename)
    
    if os.path.exists(path):
        os.remove(path)
        return {"status": "deleted", "item": item_name}
    raise HTTPException(status_code=404, detail="File not found")

# Deletes a specific message from a session file based on its index
@app.post("/api/sessions/delete-message")
def delete_message(req: DeleteMessageRequest):
    path = os.path.join(SESSIONS_DIR, req.filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Session not found")
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        idx = req.msg_index
        if 0 <= idx < len(data['history']):
            del data['history'][idx:idx+2]
            
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Removes a specific item from the browser history within a session file    
@app.delete("/api/browser-history/delete-item")
def delete_browser_history_item(session_file: str, index: int):
    path = os.path.join(SESSIONS_DIR, session_file)
    
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Session file not found")
        
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        if "browser_history" in data:
            history = data["browser_history"]
            
            if 0 <= index < len(history):
                removed_url = history.pop(index)
                
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=4, ensure_ascii=False)
                    
                return {"status": "success", "removed": removed_url}
            else:
                raise HTTPException(status_code=400, detail="Invalid history index")
        else:
            return {"status": "no_history_found"}
            
    except Exception as e:
        print(f"Error deleting history entry: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Returns a list of all files currently stored in the knowledge base directory
@app.get("/api/files")
def get_files():
    if not os.path.exists(KNOWLEDGE_DIR): return {"files": []}
    files = [f for f in os.listdir(KNOWLEDGE_DIR) if os.path.isfile(os.path.join(KNOWLEDGE_DIR, f))]
    return {"files": files}

# Uploads files and triggers the backend 4-4-2 Indexing system logic
@app.post("/api/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """
    Uploads files and triggers the 4-4-2 Indexing system logic.
    """
    saved_files = []
    for file in files:
        try:
            file_location = os.path.join(KNOWLEDGE_DIR, file.filename)
            with open(file_location, "wb+") as file_object:
                shutil.copyfileobj(file.file, file_object)
            saved_files.append(file.filename)
        except Exception: continue
    
    if saved_files:
        try:
            engine.index_documents()
            print(f"✅ Indexed {len(saved_files)} new files via 4-4-2 System.")
        except Exception as e:
            print(f"⚠️ Indexing error: {e}")

    return {"info": f"Saved & Indexed {len(saved_files)} files", "files": saved_files}

# Deletes a specific file from the knowledge base directory
@app.delete("/api/files/{filename}")
def delete_file(filename: str):
    path = os.path.join(KNOWLEDGE_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
        return {"status": "deleted", "file": filename}
    raise HTTPException(status_code=404, detail="File not found")

# Returns a list of all available chat sessions
@app.get("/api/sessions")
def get_sessions():
    if not os.path.exists(SESSIONS_DIR): return {"sessions": []}
    
    files = [f for f in os.listdir(SESSIONS_DIR) if f.endswith(".json") and not f.startswith("temp_passive_")]
    sessions = []
    for f in files:
        try:
            with open(os.path.join(SESSIONS_DIR, f), 'r', encoding='utf-8') as file:
                data = json.load(file)
                sessions.append({
                    "filename": f,
                    "id": data.get("id", 0),
                    "name": data.get("name", "Unknown")
                })
        except: pass
    sessions.sort(key=lambda x: x['id'], reverse=True)
    return {"sessions": sessions}

# Returns the most recent session file, including temporary ones, for synchronization
@app.get("/api/sessions/latest")
def get_latest_session():
    """Hidden sync endpoint: Returns the latest file, including temp files."""
    if not os.path.exists(SESSIONS_DIR): 
        return {"filename": None}
        
    files = [f for f in os.listdir(SESSIONS_DIR) if f.endswith(".json")]
    if not files:
        return {"filename": None}
        
    files.sort(key=lambda x: os.path.getmtime(os.path.join(SESSIONS_DIR, x)), reverse=True)
    
    return {"filename": files[0]}

# Loads and returns the history, name, ID, and browser history of a specific session
@app.get("/api/sessions/{filename}")
def load_session(filename: str):
    path = os.path.join(SESSIONS_DIR, filename)
    if not os.path.exists(path):
        return {"history": [], "name": "Unknown", "id": 0, "browser_history": []}
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return {
            "history": data.get("history", []), 
            "name": data.get("name"), 
            "id": data.get("id"),
            "browser_history": data.get("browser_history", []) 
        }
    except:
        return {"history": [], "name": "Error Loading", "browser_history": []}

# Creates a new session and reindexes existing sessions
@app.post("/api/sessions/create")
def create_session_api():
    reindex_sessions()
    
    dir_path = str(SESSIONS_DIR)
    existing_files = [f for f in os.listdir(dir_path) if f.endswith('.json')]
    new_id = len(existing_files) + 1
    filename = f"session_{new_id}.json"
    
    pretty_name = f"Session #{new_id}"
    
    data = {
        "id": new_id, 
        "name": pretty_name, 
        "created_at": str(datetime.now()), 
        "history": [], 
        "browser_history": []
    }
    
    file_path = os.path.join(dir_path, filename)
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4, ensure_ascii=False)
        
    return {"filename": filename, "name": pretty_name}

# Renames an existing session file and updates its internal metadata
@app.post("/api/sessions/rename")
def rename_session(req: RenameRequest):
    path = os.path.join(SESSIONS_DIR, req.filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Session file not found")
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        data['name'] = req.new_name
        
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
            f.flush() 
            os.fsync(f.fileno()) 
        
        if "reindex_sessions" in globals():
            reindex_sessions()
            
        return {"status": "success", "new_name": req.new_name}
    except Exception as e:
        print(f"Rename error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Deletes a specified session file and triggers a reindexing of the remaining sessions
@app.delete("/api/sessions/{filename}")
def delete_session(filename: str):
    path = os.path.join(SESSIONS_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
        reindex_sessions()
        return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="Session not found")

# Clears the browser history stored within a specific session file
@app.post("/api/browser-history/clear")
async def clear_history(session_file: str):
    file_path = os.path.join(SESSIONS_DIR, session_file)
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        data['browser_history'] = [] 
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4)
        return {"status": "success"}
    return {"status": "error", "message": "File not found"}

# Executes the Ollama CLI to retrieve and return a list of available local LLMs
@app.get("/api/ollama-models")
def get_ollama_models():
    ollama_url = os.getenv("OLLAMA_HOST", "http://localhost:11434")
    try:
        response = requests.get(f"{ollama_url}/api/tags", timeout=5)
        response.raise_for_status()
        data = response.json()
        
        models = [model["name"] for model in data.get("models", [])]
        return {"models": models}
    except Exception as e:
        print(f"Error connecting to Ollama container: {e}")
        return {"models": []}

# Main chat endpoint supporting AGI modes, instruction priming, and hybrid context evaluation
@app.post("/chat")
async def chat(req: ChatRequest, request: Request):
    """
    Main Chat Endpoint supporting AGI modes, Instruction Priming, and Hybrid Context.
    """
    personality_sys = load_sys_content("personalities", req.personality)
    emotion_sys = load_sys_content("emotionalities", req.emotion)
    
    values_sys = ""
    for val_name in req.values:
        values_sys += load_sys_content("values", val_name) + "\n"

    language_rule = "\n\nRespond in the language used by the user in the last prompt."

    profile_block = f"""
### COGNITIVE PROFILE:
PERSONALITY: {personality_sys}
EMOTION: {emotion_sys}
VALUES: {values_sys}
{language_rule}
"""

    session_file_to_use = req.session_file
    
    if req.mode in [1, 2] and not session_file_to_use:
        temp_filename = f"temp_passive_{uuid.uuid4().hex[:8]}.json"
        session_file_to_use = temp_filename
        file_path = os.path.join(SESSIONS_DIR, temp_filename)
        TEMP_PASSIVE_SESSIONS.add(file_path) 
        
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump({
                "id": int(datetime.now().timestamp()), 
                "name": "Passive Session",
                "created_at": str(datetime.now()),
                "history": [],
                "browser_history": []
            }, f, ensure_ascii=False, indent=4)

    actual_api_key = load_sys_content("serper_keys", req.serper_key) if req.web_search else None
    
    stream_generation_func = engine.get_chain(
        model_name=req.model, 
        mode_id=req.mode, 
        current_session_file=session_file_to_use, 
        combined_files=req.combined_sessions, 
        profile_config=profile_block,
        selected_files=req.selected_files,
        web_search_enabled=req.web_search,
        serper_key=actual_api_key,
        insight_enabled=req.insight,        
        temperature=req.temperature,
        browser_context=req.browser_context,    
        system_layers=req.system_layers      
    )
    
    async def event_generator():
        for chunk in stream_generation_func(req.message):
            if await request.is_disconnected():
                print("[SERVER] Client disconnected. Stopping generation immediately.")
                break
                
            yield chunk
            await asyncio.sleep(0)
            
    return StreamingResponse(event_generator(), media_type="text/plain")

# Serves the main index.html file for the frontend application
@app.get("/")
async def root():
    return File(os.path.join("static", "index.html"))

VIEWPORT_WIDTH = 1280
VIEWPORT_HEIGHT = 720

# Establishes a WebSocket connection for streaming browser data
@app.websocket("/ws/browser")
async def browser_stream(websocket: WebSocket):
    await websocket.accept()
    context = None
    
    try:
        # Initializes and manages the async Playwright instance context for browser automation
        async with async_playwright() as p:
            context = await p.chromium.launch_persistent_context(
                user_data_dir,
                headless=True,
                viewport={"width": VIEWPORT_WIDTH, "height": VIEWPORT_HEIGHT},
                device_scale_factor=2.0,
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
                args=[
                    "--disable-blink-features=AutomationControlled", 
                    "--no-sandbox",
                    "--disable-dev-shm-usage"
                ]
            )
            
            page = context.pages[0] if context.pages else await context.new_page()
            await page.add_init_script("Object.defineProperty(navigator, 'webdriver', { get: () => undefined });")

            cdp = await page.context.new_cdp_session(page)

            is_stream_active = True 

            # Stops any existing screencast session and initiates a new one to capture visual browser frames
            async def restart_screencast():
                try: await cdp.send("Page.stopScreencast")
                except Exception: pass
                
                if is_stream_active:
                    try:
                        await cdp.send("Page.startScreencast", {"format": "jpeg", "quality": 80, "everyNthFrame": 3})
                        await page.evaluate("window.scrollBy(0, 1); window.scrollBy(0, -1);")
                    except Exception: pass

            is_sending_frame = False

            # Event handler triggered when a new screencast frame is emitted, typically responsible for processing and sending the frame data over the WebSocket
            async def on_screencast_frame(event):
                nonlocal is_sending_frame
                
                if is_sending_frame:
                    try: await cdp.send("Page.screencastFrameAck", {"sessionId": event["sessionId"]})
                    except Exception: pass
                    return

                is_sending_frame = True
                try:
                    await websocket.send_json({"type": "frame", "data": event["data"]})
                except Exception: 
                    pass
                finally:
                    try: await cdp.send("Page.screencastFrameAck", {"sessionId": event["sessionId"]})
                    except Exception: pass
                    is_sending_frame = False

            cdp.on("Page.screencastFrame", lambda ev: asyncio.create_task(on_screencast_frame(ev)))

            # Event handler triggered upon browser navigation to track, log, or process page transitions within the virtual browser
            async def on_nav(frame):
                if frame == page.main_frame:
                    try: await websocket.send_json({"type": "url_changed", "url": frame.url})
                    except Exception: pass
            
            page.on("framenavigated", lambda fr: asyncio.create_task(on_nav(fr)))

            # MAIN EVENT LOOP: ASYNCHRONOUS BIDIRECTIONAL BROWSER AUTOMATION VIA WS
            while True:
                try:
                    # Asynchronously block and listen for inbound JSON payloads from the frontend client
                    msg = await websocket.receive_json()
                except Exception:
                    # Break the loop and trigger cleanup mechanisms if the client disconnects or connection drops
                    break 

                action = msg.get("action")

                # BROWSER RESILIENCE LAYER: SELF-HEALING / CRASH REVIVAL MECHANISM
                if page.is_closed():
                    print("The virtual browser crashed. Revival...")
                    try:
                        # Prevent memory leaks by detaching event listeners from the dead CDP session
                        try: cdp.remove_all_listeners("Page.screencastFrame")
                        except Exception: pass
                        
                        # Spawn a fresh browser tab/page instance within the persistent state context
                        page = await context.new_page()

                        # Establish a new Chrome DevTools Protocol (CDP) session for the revived tab
                        cdp = await page.context.new_cdp_session(page)
                        
                        # Re-attach asynchronous event listeners for screencasting and navigation telemetry
                        cdp.on("Page.screencastFrame", lambda ev: asyncio.create_task(on_screencast_frame(ev)))
                        page.on("framenavigated", lambda fr: asyncio.create_task(on_nav(fr)))
                        
                        # Synchronize screen capturing with the newly generated page state
                        await restart_screencast()
                    except Exception as e:
                        print(f"Critical crash of Context: {e}")

                # Re-evaluate the payload action after potential execution of the revival sequence
                action = msg.get("action")

                # STREAM CONTROL INTERFACES: RESOURCE OPTIMIZATION
                if action == "pause_stream":
                    # Global flag interception to block execution of automated screencast triggers
                    is_stream_active = False
                    try: await cdp.send("Page.stopScreencast")
                    except Exception: pass
                    continue

                elif action == "resume_stream":
                    # Re-activate global streaming state and force-initialize the visual pipeline
                    is_stream_active = True
                    await restart_screencast()
                    continue

                # EXECUTION LAYER: AGENTIC INPUT TRANSLATION & INJECTION
                try:
                    # ACTION: GOTO (Target URL Navigation & Query Parsing)
                    if action == "goto":
                        url = msg.get("url", "").strip()
                        
                        # Guard clause against void inputs, empty states, or unsafe inline protocols
                        if not url or url == "about:blank" or url.startswith("data:"):
                            continue

                        # Simple utility helper function intended to normalize or sanitize inputs, such as URLs or resource paths    
                        def normalize(u): return u.replace("https://", "").replace("http://", "").replace("www.", "").rstrip("/")
                        
                        # Prevent duplicate network overhead if the requested URL matches the current active state
                        if not page.url.startswith("data:") and normalize(page.url) == normalize(url):
                            await restart_screencast()
                            continue

                        # Regex assertion to evaluate if input is a formatted URL or a standard raw string query    
                        is_valid_url = re.match(r'^(https?://)?[a-z0-9-]+(\.[a-z0-9-]+)+\.*', url, re.IGNORECASE)
                        has_space = " " in url

                        # Transform raw string entries or structurally invalid inputs into Google Search pathways
                        if has_space or not is_valid_url:
                            query = urllib.parse.quote(url.replace("https://", "").replace("http://", ""))
                            url = f"https://www.google.com/search?q={query}"
                        elif not url.startswith(("http://", "https://")):
                            url = "https://" + url

                        try:
                            # Execute navigation optimizing for DOM availability to minimize process latency
                            await page.goto(url, wait_until="domcontentloaded", timeout=15000)
                            await restart_screencast()
                        except Exception as e:
                            err_str = str(e)
                            # Fail-safe routing: fallback to index searching upon DNS or network resolution failures
                            if "ERR_NAME_NOT_RESOLVED" in err_str or "ERR_INVALID_URL" in err_str:
                                fallback_query = urllib.parse.quote(url.replace("https://", "").replace("http://", ""))
                                fallback_url = f"https://www.google.com/search?q={fallback_query}"
                                try: 
                                    await page.goto(fallback_url, wait_until="domcontentloaded", timeout=10000)
                                    await restart_screencast()
                                except Exception: pass

                    # --- ACTION: MOUSEMOUSE (Simulated Coordinate Mapping)
                    elif action == "mousemove":
                        try: await page.mouse.move(msg["x"], msg["y"])
                        except Exception: pass
                    
                    # --- ACTION: MOUSEDOWN (Simulated Primitive Pointer Engagements)
                    elif action == "mousedown":
                        try: await page.mouse.down()
                        except Exception: pass
                    
                    # ACTION: MOUSEUP (Simulated Primitive Pointer Disengagements)
                    elif action == "mouseup":
                        try: await page.mouse.up()
                        except Exception: pass
                    
                    # ACTION: CLICK (Discreet Vector Element Interaction)
                    elif action == "click":
                        try: await page.mouse.click(msg["x"], msg["y"])
                        except Exception: pass
                    
                    # ACTION: SCROLL (Simulated Asynchronous Desktop Wheel Grids)
                    elif action == "scroll":
                        try: await page.mouse.wheel(0, msg["deltaY"])
                        except Exception: pass
                
                    # ACTION: SHORTCUT (System Keyboard Macros Execution)
                    elif action == "shortcut":
                        try: await page.keyboard.press(msg["key"])
                        except Exception: pass
                    
                    # ACTION: PASTE (Direct Document Object Buffer Injection)
                    elif action == "paste":
                        try: await page.keyboard.insert_text(msg["text"])
                        except Exception: pass
                    
                    # ACTION: KEYDOWN (Advanced Hardware Keyboard Event emulation)
                    elif action == "keydown":
                        try:
                            key = msg["key"]
                            if key in ["Control", "Shift", "Alt", "Meta", "AltGraph"]: continue
                            if len(key) == 1: await page.keyboard.insert_text(key)
                            else: await page.keyboard.press(key)
                        except Exception: pass
                    
                    # ACTION: GET_SELECTION (Targeted Semantic Extraction for AI Context Layers)
                    elif action == "get_selection":
                        try:
                            # Query the active DOM selection bounding boxes asynchronously
                            selected_text = await page.evaluate("window.getSelection().toString().trim()")
                            if selected_text:
                                # Capture full inner text metrics to supply structural context around the subset extraction
                                page_text = await page.evaluate("document.body.innerText")
                                await websocket.send_json({"type": "selection_result", "text": selected_text, "rect": msg.get("rect"), "page_context": page_text})
                            else:
                                # Echo null structures back to frontend if selection yields no actual characters
                                await websocket.send_json({"type": "selection_result", "text": "", "rect": None, "page_context": ""})
                        except Exception: pass
                    
                    # ACTION: GO_BACK (Browser Session History Decrement)
                    elif action == "go_back":
                        try: await page.go_back(wait_until="domcontentloaded", timeout=5000)
                        except Exception: pass

                    # --- ACTION: GO_FORWARD (Browser Session History Increment)
                    elif action == "go_forward":
                        try: await page.go_forward(wait_until="domcontentloaded", timeout=5000)
                        except Exception: pass

                    # ACTION: GET_FULL_CONTEXT (Deep Extraction Pipe feeding the AGI / Samsara Layer)
                    elif action == "get_full_context":
                        try:
                            page_text = await page.evaluate("document.body ? document.body.innerText : ''")
                            await websocket.send_json({"type": "full_context_result", "text": page_text})
                        except Exception: pass

                    # ACTION: EXPORT_PNG (On-Demand Visual State Frame Capture)
                    elif action == "export_png":
                        export_type = msg.get("type")
                        try:
                            screenshot_bytes = b""
                            # Route screenshot strategies based on context-specific parameters
                            if export_type == "visible": screenshot_bytes = await page.screenshot(type="png")
                            elif export_type == "full": screenshot_bytes = await page.screenshot(type="png", full_page=True)
                            elif export_type == "selected" and msg.get("clip"): screenshot_bytes = await page.screenshot(type="png", clip=msg.get("clip"))

                            if screenshot_bytes:
                                # Encode output binary stream to standardized Base64 ASCII for over-the-network safely
                                b64_data = base64.b64encode(screenshot_bytes).decode('utf-8')
                                current_url = page.url
                                domain = "magi_screenshot"
                                try: domain = urllib.parse.urlparse(current_url).netloc.replace("www.", "")
                                    # Parse target domain endpoints to generate clean file semantic names
                                except Exception: pass
                                await websocket.send_json({"type": "export_result", "data": b64_data, "filename": f"{domain}_{export_type}.png"})
                        except Exception: pass

                except Exception:
                    # Catch and isolate internal handler exceptions to maintain main loop structural resilience
                    pass 

    except asyncio.CancelledError:
        print("The server is shutting down..")
    except Exception as e:
        print(f"Critical error in virtual browser: {e}")
    finally:
        print("Clear browser...")
        if 'cdp' in locals() and cdp:
            try: await cdp.send("Page.stopScreencast")
            except Exception: pass
        if 'context' in locals() and context:
            try: await context.close()
            except Exception: pass

# Adds a new URL entry to the browser history of a specified session
@app.post("/api/sessions/add-browser-history")
def add_browser_history(data: dict):
    filename = data.get("filename")
    url = data.get("url")
    
    if not filename or not url or url == "about:blank":
        return {"status": "ignored"}

    path = os.path.join(SESSIONS_DIR, filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Session not found")

    try:
        with open(path, 'r', encoding='utf-8') as f:
            session_data = json.load(f)
        
        if "browser_history" not in session_data:
            session_data["browser_history"] = []
            
        history = session_data["browser_history"]
        
        if not history or history[-1] != url:
            history.append(url)
            session_data["browser_history"] = history[-50:]

            with open(path, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, indent=4, ensure_ascii=False)
        
        return {"status": "success"}
    except Exception as e:
        print(f"Error while saving history: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Modifies the text of an existing message at a specific index within a session
@app.post("/api/sessions/update-message")
async def update_message(data: dict):
    filename = data.get("filename")
    index = data.get("index")
    new_text = data.get("new_text")
    
    file_path = os.path.join(SESSIONS_DIR, filename)
    if not os.path.exists(file_path):
        return {"error": "File not found"}
        
    with open(file_path, 'r', encoding='utf-8') as f:
        session_data = json.load(f)
        
    if 0 <= index < len(session_data["history"]):
        original_prompt = session_data["history"][index]["content"]
        
        session_data["history"][index]["content"] = new_text
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(session_data, f, indent=4, ensure_ascii=False)

        ai_response_text = ""
        if index + 1 < len(session_data["history"]):
            ai_response_text = session_data["history"][index+1]["content"]
            
        engine.edit_samsara_record(original_prompt, new_text, ai_response_text)

    if filename:
        is_passive = filename.startswith("temp_passive_")
        mode_id = 2 if is_passive else 4
        engine.rebuild_dynamic_layer_vectors(mode_id, filename)

    return {"status": "ok"}

import re

# Handles inline text editing requests by calling an engine chain for stream generation
@app.post("/api/chat/inline")
async def chat_inline(req: InlineChatRequest, request: Request):
    action_type = req.type 
    target = req.target_text
    instruction = req.instruction
    model_name = req.model
    mode_id = req.mode 
    session_file = req.session_file
    
    history_context = ""
    if session_file:
        file_path = os.path.join(SESSIONS_DIR, session_file)
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
                recent_msgs = session_data.get("history", [])[-6:]
                for msg in recent_msgs:
                    role = msg.get("role", "unknown").upper()
                    content = msg.get("content", "")
                    history_context += f"{role}: {content}\n"

    prompts = []
    
    before_ctx = req.text_before[-1500:] if req.text_before else ""
    after_ctx = req.text_after[:1500] if req.text_after else ""
    
    context_block = ""
    if before_ctx or after_ctx:
        context_block = f"""
                        --- SURROUNDING CONTEXT (For style & flow adaptation ONLY. DO NOT include this in your output!) ---
                        [BEFORE]: {before_ctx}
                        [AFTER]: {after_ctx}
                        ---------------------------------------------------------------------------------------------------
                        """
    
    if action_type == "regen":
        prompts.append(f"Rewrite ONLY the following TARGET TEXT based on the instruction: '{instruction}'.\n{context_block}\n[TARGET TEXT TO REWRITE]:\n{target}")
    else: 
        prompts.append(f"Expand ONLY on the following TARGET TEXT in detail based on the overall instruction: '{instruction}'.\n{context_block}\n[TARGET TEXT TO EXPAND]:\n{target}")

    system_profile = f"""
                        You are an expert inline text editor. You are editing a specific part of a document.
                        PREVIOUS CHAT CONTEXT:
                        {history_context}

                        RULES:
                        1. Return ONLY the edited/expanded TARGET TEXT.
                        2. CRITICAL: Do NOT output or repeat the "SURROUNDING CONTEXT" (Before/After text). Your output will be directly injected into the middle of the document!
                        3. Do NOT add introductions like 'Here is the text:' or 'Sure!'.
                        4. Stop generating immediately after completing the request. Do NOT loop or add reasoning traces.
                        """

    # Asynchronously streams generated text events based on prompt configurations and system layers
    async def multi_event_generator():
        for i, p in enumerate(prompts):
            stream_generation_func = engine.get_chain(
                model_name=req.model, 
                mode_id=req.mode, 
                current_session_file=req.session_file, 
                combined_files=[], 
                profile_config=system_profile,
                selected_files=[],
                web_search_enabled=False,
                serper_key=None,
                insight_enabled=False,        
                temperature=0.7,
                system_layers=req.system_layers 
            )
            
            for chunk in stream_generation_func(p):
                if await request.is_disconnected():
                    print("🚫 [SERVER] Client disconnected during inline generation.")
                    return
                yield chunk
                await asyncio.sleep(0)
            
            if action_type == "expand" and i < len(prompts) - 1:
                yield "\n\n"
                
    return StreamingResponse(multi_event_generator(), media_type="text/plain")

# Lists files in a specified file system directory, excluding certain hidden directories
@app.get("/api/fs/list")
def list_fs_files(path: str = 'Root'):
    EXCLUDED_DIRS = {'astral_projection', 'clipboard', '__pycache__', '.git'}
    path = normalize_path(path)
    if '..' in path:
        raise HTTPException(status_code=403, detail="Invalid path")

    base_dir = BASE_DIR
    items = []

    try:
        if path == 'Root':
            for folder in ALLOWED_ROOT_FOLDERS:
                folder_path = os.path.join(base_dir, folder)
                if os.path.exists(folder_path) and os.path.isdir(folder_path):
                    items.append({
                        "name": folder,
                        "type": "folder",
                        "path": folder,
                        "immutable": True 
                    })
            
            for item in os.listdir(base_dir):
                item_path = os.path.join(base_dir, item)
                if os.path.isfile(item_path):
                    items.append({
                        "name": item,
                        "type": "file",
                        "path": item,
                        "size": os.path.getsize(item_path),
                        "immutable": False
                    })
        else:
            target_dir = os.path.join(base_dir, path)
            if not os.path.exists(target_dir) or not os.path.isdir(target_dir):
                raise HTTPException(status_code=404, detail="Directory not found")

            for item in os.listdir(target_dir):
                
                if item in ["documents.json", "ollama_list.json", "0_run_main.bat", "init_logic_db.py.py", "logic.py", "main.py"] or item in EXCLUDED_DIRS:
                    continue

                item_path = os.path.join(target_dir, item)
                is_dir = os.path.isdir(item_path)
                items.append({
                    "name": item,
                    "type": "folder" if is_dir else "file",
                    "path": f"{path}/{item}".strip("/"),
                    "size": os.path.getsize(item_path) if not is_dir else 0,
                    "immutable": False
                })

        def natural_sort_key(item):
            parts = [int(text) if text.isdigit() else text for text in re.split(r'(\d+)', item['name'].lower())]
            
            return (item['type'] == 'file', parts)

        items.sort(key=natural_sort_key)
        return items

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
class FSRenameRequest(BaseModel):
    old_path: Optional[str] = None 
    old_name: Optional[str] = None 
    new_name: str
    mode: Optional[str] = None 

# Renames a file or directory in the file system after validating paths
@app.post("/api/fs/rename")
def fs_rename(req: FSRenameRequest):
    req.old_path = normalize_path(req.old_path) 
    req.old_name = normalize_path(req.old_name) 
    req.new_name = normalize_path(req.new_name)

    if ('..' in (req.old_path or "") or '..' in (req.old_name or "") 
        or '..' in req.new_name or '/' in req.new_name):
        raise HTTPException(status_code=403, detail="Invalid path or name")

    if req.mode == 'page' and req.old_name:
        old_full_path = os.path.join(PLANCHETTE_DIR, req.old_name)

        display_name = req.new_name.replace(".jdoc.json", "").strip()

        safe_new_name = re.sub(r'\s+', '_', display_name) + ".jdoc.json"
        new_full_path = os.path.join(PLANCHETTE_DIR, safe_new_name)

        if not os.path.exists(old_full_path):
            raise HTTPException(status_code=404, detail="File not found")

        try:
            os.replace(old_full_path, new_full_path)

            with open(new_full_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            if "metadata" not in data:
                data["metadata"] = {}

            data["metadata"]["name"] = display_name 
            data["metadata"]["last_modified"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            with open(new_full_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)

            return {
                "status": "success", 
                "new_path": new_full_path, 
                "new_name": safe_new_name,     
                "display_name": display_name    
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
            
    elif req.old_path:
        old_full_path = os.path.join(BASE_DIR, req.old_path)
        new_full_path = os.path.join(os.path.dirname(old_full_path), req.new_name)
        if not os.path.exists(old_full_path):
            raise HTTPException(status_code=404, detail="File not found")
        try:
            os.rename(old_full_path, new_full_path)
            return {"status": "success", "new_path": new_full_path, "new_name": req.new_name}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    else:
        raise HTTPException(status_code=400, detail="Invalid request parameters")

# Deletes a specified path from the file system
@app.delete("/api/fs/delete")
def fs_delete(path: str):
    path = normalize_path(path)
    if '..' in path:
        raise HTTPException(status_code=403, detail="Invalid path")
        
    full_path = os.path.join(BASE_DIR, path)
    
    target_name = os.path.basename(full_path)
    parent_dir = os.path.dirname(full_path)
    if target_name in ALLOWED_ROOT_FOLDERS and parent_dir == BASE_DIR:
        raise HTTPException(status_code=403, detail="Cannot delete core system folders")
        
    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="File not found")
        
    try:
        if os.path.isdir(full_path):
            import shutil
            shutil.rmtree(full_path) 
        else:
            os.remove(full_path) 
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
# Uploads files to a specified directory path within the file system
@app.post("/api/fs/upload")
async def fs_upload(path: str = Form(...), files: List[UploadFile] = File(...)):
    if '..' in path:
        raise HTTPException(status_code=403, detail="Invalid path")
        
    target_dir = os.path.join(BASE_DIR, path)
    if not os.path.exists(target_dir) or not os.path.isdir(target_dir):
        raise HTTPException(status_code=404, detail="Directory not found")
        
    saved_files = []
    try:
        for file in files:
            file_path = os.path.join(target_dir, file.filename)
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            saved_files.append(file.filename)
            
        return {"status": "success", "files": saved_files}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Facilitates the downloading of a file from the file system, handling background task
@app.get("/api/fs/download")
def fs_download(path: str, background_tasks: BackgroundTasks): 
    path = normalize_path(path)
    if '..' in path:
        raise HTTPException(status_code=403, detail="Invalid path")
        
    full_path = os.path.join(BASE_DIR, path)
    
    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="File or folder not found")
        
    if os.path.isdir(full_path):
        tmp_dir = tempfile.mkdtemp()
        zip_filename = os.path.basename(full_path) if path != 'Root' else 'Magi_Root'
        zip_path = os.path.join(tmp_dir, zip_filename)
        
        import shutil
        shutil.make_archive(zip_path, 'zip', full_path)
        
        background_tasks.add_task(shutil.rmtree, tmp_dir, ignore_errors=True)
        
        return FileResponse(
            path=zip_path + ".zip", 
            filename=zip_filename + ".zip",
            media_type='application/zip'
        )
        
    return FileResponse(path=full_path, filename=os.path.basename(full_path))

class FSPasteRequest(BaseModel):
    action: str
    sources: List[str]
    destination: str

# Pastes, copies or moves sources to a destination directory within the file system
@app.post("/api/fs/paste")
def fs_paste(req: FSPasteRequest):
    req.destination = normalize_path(req.destination)
    if '..' in req.destination:
        raise HTTPException(status_code=403, detail="Invalid destination")
    
    dest_dir = os.path.join(BASE_DIR, req.destination)
    if not os.path.exists(dest_dir) or not os.path.isdir(dest_dir):
        raise HTTPException(status_code=404, detail="Destination directory not found")

    results = []
    for src in req.sources:
        src = normalize_path(src)
        if '..' in src: continue
        src_path = os.path.join(BASE_DIR, src)
        if not os.path.exists(src_path): continue
        
        if req.action == 'cut':
            target_name = os.path.basename(src_path)
            parent_dir = os.path.dirname(src_path)
            if target_name in ALLOWED_ROOT_FOLDERS and parent_dir == BASE_DIR:
                continue 

        dest_path = os.path.join(dest_dir, os.path.basename(src_path))
        
        if os.path.exists(dest_path) and req.action == 'copy':
            base, ext = os.path.splitext(dest_path)
            dest_path = f"{base}_copy{ext}"

        try:
            if req.action == 'copy':
                if os.path.isdir(src_path):
                    shutil.copytree(src_path, dest_path, dirs_exist_ok=True)
                else:
                    shutil.copy2(src_path, dest_path)
            elif req.action == 'cut':
                shutil.move(src_path, dest_path)
            results.append(dest_path)
        except Exception as e:
            print(f"Error inserting {src}: {e}")
            
    return {"status": "success", "pasted": len(results)}

import os

# Saves provided JSON data to the clipboard history file
@app.post("/api/clipboard/save")
async def save_to_clipboard(request: Request):
    try:
        data = await request.json()

        if not os.path.exists(CLIPBOARD_FILE):
            with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
                json.dump([], f)
                
        with open(CLIPBOARD_FILE, "r", encoding="utf-8") as f:
            history = json.load(f)
        
        new_entry = {
            "id": str(datetime.now().timestamp()),
            "timestamp": datetime.now().isoformat(),
            "action": data.get("action"),  
            "content": data.get("content")
        }
        history.append(new_entry)
        
        if len(history) > 50:
            history = history[-50:]
        
        with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=4, ensure_ascii=False)
            
        return {"status": "ok"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# Retrieves and returns the current clipboard history from its JSON storage
@app.get("/api/clipboard/history")
async def get_clipboard_history():
    try:
        if not os.path.exists(CLIPBOARD_FILE):
            return []
        with open(CLIPBOARD_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        return []

# Removes a specific item from the clipboard history based on its unique ID
@app.delete("/api/clipboard/delete/{item_id}")
async def delete_clipboard_item(item_id: str):
    try:
        with open(CLIPBOARD_FILE, "r", encoding="utf-8") as f:
            history = json.load(f)
        
        history = [item for item in history if item.get("id") != item_id]
        
        with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=4, ensure_ascii=False)
            
        return {"status": "success"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# Completely clears all items from the clipboard history file
@app.delete("/api/clipboard/clear")
async def clear_clipboard_history():
    try:
        with open(CLIPBOARD_FILE, "w", encoding="utf-8") as f:
            json.dump([], f, indent=4, ensure_ascii=False)
        return {"status": "success"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# Regenerates an AI response based on a specific prompt, context, and selected model
@app.post("/api/ai/regen")
async def ai_regen_endpoint(request: Request):
    try:
        data = await request.json()
        prompt = data.get("prompt", "").strip()
        context = data.get("context", "").strip()
        model_name = data.get("model", "llama3") 

        system_instruction = """
        You are a highly capable and direct writing assistant. 
        Your task is to follow the user's instruction precisely.
        Output ONLY the raw final text. Do not include introductory phrases, conversational filler, or internal thoughts.
        """
        
        if context:
            user_message = f"Modify the following text based on this instruction: {prompt}\n\nTEXT TO MODIFY:\n{context}"
        else:
            user_message = f"Instruction: {prompt}"
            
        messages = [
            ("system", system_instruction),
            ("human", user_message)
        ]
        
        llm = engine.get_llm(model_name, temperature=0.7)
        result = llm.invoke(messages)
        ai_text = result.content if hasattr(result, 'content') else str(result)
        
        ai_text = re.sub(r'<think>.*?</think>', '', ai_text, flags=re.DOTALL)
        ai_text = re.sub(r'^(Here is|Sure|Certainly|Here\'s|Okay|Absolutely).*?:?\s*\n*', '', ai_text.strip(), flags=re.IGNORECASE)
        
        return {"status": "success", "text": ai_text.strip()}
        
    except Exception as e:
        print(f"Regen API Error: {e}")
        return {"status": "error", "message": str(e)}

# Retrieves and returns the stored astral projection cell data
@app.get("/api/astral")
async def get_astral_data():
    try:
        with open(ASTRAL_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        return {"cells": []}

# Saves the astral projection cells data directly to the corresponding JSON file
@app.post("/api/astral/save")
async def save_astral_cells(request: Request):
    try:
        data = await request.json()
        with open(ASTRAL_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        return {"status": "ok"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

class CreateDocRequest(BaseModel):
    filename: str

# Creates a new docstral document by sanitizing the requested filename
@app.post("/api/docstral/create")
async def create_jdoc(req: CreateDocRequest):
    raw_name = req.filename.strip()
    
    base_fs_name = re.sub(r'\s+', '_', raw_name)
    if base_fs_name.endswith(".jdoc.json"):
        base_fs_name = base_fs_name[:-10]

    clean_filename = f"{base_fs_name}.jdoc.json"
    filepath = os.path.join(PLANCHETTE_DIR, clean_filename)
    
    counter = 1
    while os.path.exists(filepath):
        clean_filename = f"{base_fs_name}_{counter}.jdoc.json"
        filepath = os.path.join(PLANCHETTE_DIR, clean_filename)
        counter += 1
    
    doc_id = f"doc-{int(datetime.now().timestamp())}"
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    empty_doc = {
        "file_version": "1",
        "metadata": {
            "id": doc_id,
            "name": raw_name,
            "created_at": now_str,
            "last_modified": now_str               
        },

        "visible_layers": ['title', 'h1_names', 'h2_names', 'h3_names', 'h4_names', 'h1_content', 'h2_content', 'h3_content', 'h4_content', 'content', 'quotes'],
        "blocks": [
            {
                "id": generate_new_block_id(),
                "type": "paragraph",
                "content": "<br>",
                "layer": "base"
            }
        ]
    }
    
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(empty_doc, f, indent=4, ensure_ascii=False)
            
        return {
            "status": "success", 
            "filename": clean_filename, 
            "data": empty_doc
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Returns a list of files filtered and sorted depending on the active Docstral mode
@app.get("/api/fs/list_by_mode")
def list_fs_by_mode(mode: str = Query(...)):
    """
    Returns a list of files depending on Docstral mode.
    Sorted by date (newest documents at the top).
    """
    try:
        if mode == 'page':
            if not os.path.exists(PLANCHETTE_DIR):
                return {"files": []}
            
            files_with_dates = []
            
            for filename in os.listdir(PLANCHETTE_DIR):
                if filename.endswith('.jdoc.json') or filename.endswith('.json'):
                    filepath = os.path.join(PLANCHETTE_DIR, filename)
                    file_date = None
                    
                    try:
                        with open(filepath, "r", encoding="utf-8") as f:
                            data = json.load(f)
                            meta = data.get("metadata", {})
                            file_date = meta.get("created_at") or meta.get("last_modified")
                    except Exception:
                        pass 
                        
                    if not file_date:
                        file_date = datetime.fromtimestamp(os.path.getctime(filepath)).strftime("%Y-%m-%d %H:%M:%S")
                        
                    files_with_dates.append({
                        "filename": filename,
                        "date": file_date
                    })
            
            files_with_dates.sort(key=lambda x: str(x["date"]), reverse=True)
            
            sorted_files = [f["filename"] for f in files_with_dates]
            
            return {"files": sorted_files}
        
        return {"files": []}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Loads a specific document file corresponding to the requested operating mode    
@app.get("/api/fs/load_file")
def load_doc_file(mode: str, filename: str):
    """
    Loads a specific document for the corresponding mode.
    """
    if '..' in filename or '/' in filename:
        raise HTTPException(status_code=403, detail="Invalid filename")

    if mode == 'page':
        filepath = os.path.join(PLANCHETTE_DIR, filename)
        if os.path.exists(filepath):
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    doc_data = json.load(f)
                return {"id": f"doc-{uuid.uuid4().hex[:8]}", "name": filename, "data": doc_data}
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error reading file: {str(e)}")
                
    raise HTTPException(status_code=404, detail="File not found")

class DocstralBlock(BaseModel):
    id: str
    type: str
    content: Any = ""
    layer: Optional[str] = "base"

class DocstralRenderRequest(BaseModel):
    blocks: List[DocstralBlock]
    global_settings: Optional[Dict[str, Any]] = None

# Processes and returns Docstral document blocks as a unified stream without pagination
@app.post("/api/docstral/render")
async def render_docstral_document(req: DocstralRenderRequest):
    """
    Version 1: Infinite Page.
    No more pagination and complex height calculations!
    We simply accept, validate, and return blocks as a single stream.
    """
    try:
        blocks_list = [block.model_dump() for block in req.blocks]
        
        return {
            "status": "success",
            "blocks": blocks_list,
            "total_blocks": len(blocks_list)
        }
    except Exception as e:
        print(f"Docstral Render Error: {e}")
        return {"status": "error", "message": str(e)}
    
class SaveDocRequest(BaseModel):
    filename: str
    old_filename: Optional[str] = None
    data: dict

# Saves the provided document data and metadata into a jdoc JSON file
@app.post("/api/docstral/save")
async def save_jdoc(req: SaveDocRequest):
    data = req.data
    metadata = data.get("metadata", {})
    doc_id = metadata.get("id")
    display_name = metadata.get("name", "Untitled")
    
    active_version = metadata.get("active_version", "main")
    
    target_filename = re.sub(r'\s+', '_', display_name.strip())
    if not target_filename.endswith(".jdoc.json"):
        target_filename += ".jdoc.json"
        
    target_path = os.path.join(PLANCHETTE_DIR, target_filename)

    try:
        found_file_path = None

        if doc_id:
            for filename in os.listdir(PLANCHETTE_DIR):
                if filename.endswith(".jdoc.json"):
                    path = os.path.join(PLANCHETTE_DIR, filename)
                    try:
                        with open(path, "r", encoding="utf-8") as f:
                            existing_data = json.load(f)
                            if existing_data.get("metadata", {}).get("id") == doc_id:
                                found_file_path = path
                                break
                    except Exception:
                        continue

        if found_file_path and found_file_path != target_path:
            os.replace(found_file_path, target_path)
            existing_path = target_path
        elif found_file_path:
            existing_path = found_file_path
        else:
            existing_path = target_path 

        if os.path.exists(existing_path):
            os_ctime = os.path.getctime(existing_path)
            created_at_str = datetime.fromtimestamp(os_ctime).strftime("%Y-%m-%d %H:%M:%S")
        else:
            created_at_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        data["metadata"]["created_at"] = metadata.get("created_at") or created_at_str
        data["metadata"]["last_modified"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        data["metadata"].pop("size_kb", None)
        data["metadata"].pop("size_bytes", None)

        old_versions = {}
        old_main_blocks = []
        old_main_visible_layers = []
        
        if os.path.exists(existing_path):
            try:
                with open(existing_path, "r", encoding="utf-8") as f:
                    old_data = json.load(f)
                    old_versions = old_data.get("versions", {})
                    old_main_blocks = old_data.get("blocks", [])
                    old_main_visible_layers = old_data.get("visible_layers", [])
            except Exception:
                pass

        incoming_blocks = data.get("blocks", [])
        incoming_visible_layers = data.get("visible_layers")
        if incoming_visible_layers is None:
            incoming_visible_layers = ['title', 'h1_names', 'h2_names', 'h3_names', 'h4_names', 'h1_content', 'h2_content', 'h3_content', 'h4_content', 'content', 'quotes']

        if active_version == "main":
            data["versions"] = old_versions
            data["visible_layers"] = incoming_visible_layers
        else:
            data["blocks"] = old_main_blocks
            data["visible_layers"] = old_main_visible_layers if old_main_visible_layers else incoming_visible_layers
            
            v_created = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            if active_version in old_versions and "created_at" in old_versions[active_version]:
                v_created = old_versions[active_version]["created_at"]
                
            old_versions[active_version] = {
                "created_at": v_created,
                "last_modified": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "visible_layers": incoming_visible_layers,
                "blocks": incoming_blocks
            }
            data["versions"] = old_versions

        perfectly_ordered_data = {
            "file_version": data.get("file_version", "3.0"),
            "metadata": data["metadata"],
            "visible_layers": data["visible_layers"],
            "blocks": data["blocks"],
            "versions": data.get("versions", {})
        }

        with open(target_path, "w", encoding="utf-8") as f:
            json.dump(perfectly_ordered_data, f, indent=4, ensure_ascii=False)
            
        return {"status": "success", "actual_filename": target_filename}

    except Exception as e:
        return {"status": "error", "message": str(e)}
    
class DeleteFileRequest(BaseModel):
    filename: str
    mode: str

# Deletes a file specifically related to the docstral interface operating mode
@app.post("/api/fs/delete")
async def docstral_delete_file(req: DeleteFileRequest):
    if req.mode != 'page':
        return {"status": "error", "message": "Invalid mode for this operation."}

    if '..' in req.filename or '/' in req.filename or '\\' in req.filename:
        return {"status": "error", "message": "Invalid file name."}

    filepath = os.path.join(PLANCHETTE_DIR, req.filename)
    
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            return {"status": "success"}
        else:
            return {"status": "error", "message": "The file does not exist.."}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# Generates a unique hexadecimal ID for docstral document blocks to prevent duplication
def generate_new_block_id():
    """Generates a unique ID for each block (duplicate protection)."""
    return f"b-{uuid.uuid4().hex[:12]}"

# Cleans a filename by removing dangerous characters, extensions, and standardizing spaces
def sanitize_filename(name: str) -> str:
    """Cleans the file name of dangerous characters and removes extensions."""
    if not name:
        return "imported_doc"
    name = os.path.basename(name)
    name = re.sub(r'(?i)\.(jdoc\.json|json|txt|docx)$', '', name)
    name = re.sub(r'[^\w\sа-яА-Я-]', '', name)
    name = re.sub(r'[\s]+', '_', name)
    return name.strip('_') or "imported_doc"

# Converts a Word document paragraph into an HTML representation, preserving basic text formatting
def get_paragraph_html(para):
    """Converts a paragraph from Word to HTML with preserved Bold, Italic, Underline formatting."""
    html_content = ""
    try:
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            text = text.replace('\n', ' ') 
            
            if getattr(run, "bold", False):
                text = f"<b>{text}</b>"
            if getattr(run, "italic", False):
                text = f"<i>{text}</i>"
            if getattr(run, "underline", False):
                text = f"<u>{text}</u>"
            
            html_content += text
    except Exception:
        pass 
    return html_content.strip()

# Iterates through and extracts paragraphs and tables from a DOCX file in sequential order
def get_docx_elements_in_order(doc):
    """Iterates through paragraphs and tables in the order they appear in the document."""
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    for child in doc.element.body.iterchildren():
        try:
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)
        except Exception as e:
            print(f"Skipping corrupt DOCX element: {e}")
            continue

# Handles the uploading and importing of a document file
@app.post("/api/docstral/import")
async def import_document(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing file name")
        
    filename = file.filename.strip()
    filename_lower = filename.lower()

    doc_id = f"doc-{int(datetime.now().timestamp())}"
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    safe_name = sanitize_filename(filename)
    new_filename = f"{safe_name}.jdoc.json"
    
    counter = 1
    base_name = safe_name
    while os.path.exists(os.path.join(PLANCHETTE_DIR, new_filename)):
        new_filename = f"{base_name}_{counter}.jdoc.json"
        counter += 1
        
    ui_name = new_filename.replace('.jdoc.json', '').replace('_', ' ')
    
    new_data = None
    all_blocks = []
    
    try:
        content = await file.read()
        
        if filename_lower.endswith('.json') or filename_lower.endswith('.jdoc.json'):
            try:
                text_content = content.decode('utf-8-sig')
                data = json.loads(text_content)
            except json.JSONDecodeError:
                text_content = content.decode('windows-1251', errors='replace')
                data = json.loads(text_content)

            existing_path = os.path.join(PLANCHETTE_DIR, filename)
            existing_versions = {}
            if os.path.exists(existing_path):
                try:
                    with open(existing_path, 'r', encoding='utf-8') as f:
                        old_doc = json.load(f)
                        existing_versions = old_doc.get("versions", {})
                except: pass
                
            if isinstance(data, dict) and "blocks" in data:
                new_data = data.copy()
                new_data["versions"] = new_data.get("versions", existing_versions)
                if "metadata" not in new_data:
                    new_data["metadata"] = {}
                    
                new_data["metadata"]["id"] = doc_id
                new_data["metadata"]["name"] = ui_name
                new_data["metadata"]["created_at"] = now_str
                new_data["metadata"]["last_modified"] = now_str
            else:
                raw_blocks = data.get("blocks") if isinstance(data, dict) else (data if isinstance(data, list) else [])
                for b in raw_blocks:
                    if isinstance(b, dict):
                        nb = b.copy()
                        nb['id'] = generate_new_block_id()
                        all_blocks.append(nb)
                        
        elif filename_lower.endswith('.txt'):
            try: text = content.decode('utf-8-sig')
            except UnicodeDecodeError: text = content.decode('windows-1251', errors='replace')
            
            for p in text.split('\n'):
                clean_text = p.strip()
                if clean_text:
                    all_blocks.append({"id": generate_new_block_id(), "type": "paragraph", "content": clean_text, "layer": "base"})
                    
        elif filename_lower.endswith('.docx'):
            doc = docx.Document(BytesIO(content))
            for element in get_docx_elements_in_order(doc):
                if isinstance(element, Paragraph):
                    text_html = get_paragraph_html(element)
                    if text_html:
                        style_name = getattr(element.style, "name", "normal") if element.style else "normal"
                        style_name = style_name.lower()
                        block_type = "paragraph"
                        
                        if "title" in style_name: block_type = "title"
                        elif "heading 1" in style_name: block_type = "h1"
                        elif "heading 2" in style_name: block_type = "h2"
                        elif "heading 3" in style_name: block_type = "h3"
                        elif "heading 4" in style_name: block_type = "h4"
                        elif "quote" in style_name: block_type = "quote"
                        
                        all_blocks.append({"id": generate_new_block_id(), "type": block_type, "content": text_html, "layer": "base"})
                elif isinstance(element, Table):
                    for row in element.rows:
                        row_cells_text = []
                        for cell in row.cells:
                            cell_content = " ".join([get_paragraph_html(p) for p in cell.paragraphs if p]).strip()
                            if cell_content: row_cells_text.append(cell_content)
                        if row_cells_text:
                            all_blocks.append({"id": generate_new_block_id(), "type": "paragraph", "content": " | ".join(row_cells_text), "layer": "base"})
            
        else:
            raise ValueError(f"Unsupported file format for file: {filename}")
            
    except Exception as e:
        print(f"Import Error details: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

    if new_data is None:
        if not all_blocks:
            all_blocks = [{"id": generate_new_block_id(), "type": "paragraph", "content": "<br>", "layer": "base"}]
            
        new_data = {
            "file_version": "1",
            "metadata": {
                "name": ui_name,
                "id": doc_id,
                "created_at": now_str,
                "last_modified": now_str
            },
            "visible_layers": ['title', 'h1_names', 'h2_names', 'h3_names', 'h4_names', 'h1_content', 'h2_content', 'h3_content', 'h4_content', 'content', 'quotes'],
            "versions": {},
            "blocks": all_blocks
        }
        
    os.makedirs(PLANCHETTE_DIR, exist_ok=True)
    save_path = os.path.join(PLANCHETTE_DIR, new_filename) 
    
    try:
        with open(save_path, 'w', encoding='utf-8') as f:
            json.dump(new_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Write error: {str(e)}")
        
    return {"status": "success", "filename": new_filename, "data": new_data}

# Exports the docstral render request data into a downloadable DOCX document file
@app.post("/api/docstral/export/docx")
async def export_docx(req: DocstralRenderRequest):
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for block in req.blocks:
        if block.type == 'page_break':
            doc.add_page_break()
            continue

        content = block.content
        if not content.strip() or content.strip() == '<br>':
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        if block.type == "title":
            p.style = doc.styles['Title']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block.type in ["heading-1", "h1"]: p.style = doc.styles['Heading 1']
        elif block.type in ["heading-2", "h2"]: p.style = doc.styles['Heading 2']
        elif block.type in ["heading-3", "h3"]: p.style = doc.styles['Heading 3']
        elif block.type in ["heading-4", "h4"]: p.style = doc.styles['Heading 4']
        elif block.type == "quote": p.style = doc.styles['Intense Quote']
        elif block.type in ["paragraph", "text"]: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 

        parts = re.split(r'(</?[biu]>)', content, flags=re.IGNORECASE)
        is_bold, is_italic, is_underline = False, False, False
        
        for part in parts:
            part_lower = part.lower()
            if part_lower == '<b>': is_bold = True
            elif part_lower == '</b>': is_bold = False
            elif part_lower == '<i>': is_italic = True
            elif part_lower == '</i>': is_italic = False
            elif part_lower == '<u>': is_underline = True
            elif part_lower == '</u>': is_underline = False
            elif part:
                clean_text = re.sub(r'<[^>]+>', '', part).replace('&nbsp;', ' ').replace('&amp;', '&')
                if clean_text:
                    run = p.add_run(clean_text)
                    run.bold = is_bold; run.italic = is_italic; run.underline = is_underline
                    if block.type not in ["title", "heading-1", "h1", "heading-2", "h2", "heading-3", "h3", "heading-4", "h4"]:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=exported_document.docx"}
    )

class TOCGenerateRequest(BaseModel):
    filename: str
    target_line: int
    included_types: List[str]

# Generates an interactive Table of Contents with hyperlinks without empty lines
@app.post("/api/docstral/generate_toc")
async def generate_toc(req: TOCGenerateRequest):
    """
    Generates interactive Table of Contents (without blank lines and dots, with hyperlinks).
    """
    if not req.filename or '..' in req.filename or '/' in req.filename:
        raise HTTPException(status_code=400, detail="Invalid filename")

    filepath = os.path.join(PLANCHETTE_DIR, req.filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)

        blocks = data.get("blocks", [])
        if not blocks:
            return {"status": "error", "message": "Document is empty"}

        toc_entries = []
        for index, block in enumerate(blocks):
            b_type = block.get("type", "")
            
            if b_type == 'heading-1': b_type = 'h1'
            if b_type == 'heading-2': b_type = 'h2'
            if b_type == 'heading-3': b_type = 'h3'
            if b_type == 'heading-4': b_type = 'h4'

            if b_type in req.included_types:
                raw_text = re.sub(r'<[^>]+>', '', block.get("content", "")).replace('&nbsp;', ' ').strip()
                block_id = block.get("id", "")
                
                if raw_text and block_id:
                    toc_entries.append({
                        "text": raw_text,
                        "type": b_type,
                        "target_id": block_id
                    })

        if not toc_entries:
            return {"status": "error", "message": "No headings found to index."}

        level_map = {'title': 0, 'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4}
        new_blocks = []

        def create_block(b_type, content):
            return {
                "id": generate_new_block_id(),
                "type": b_type,
                "content": content,
                "layer": "base"
            }

        new_blocks.append(create_block("paragraph", "<br>"))
        new_blocks.append(create_block("title", "Table of Contents"))

        toc_lines_html = []
        for entry in toc_entries:
            indent_px = level_map.get(entry["type"], 0) * 20
            
            link_html = f"""<span style="display: inline-block; padding-left: {indent_px}px; line-height: 1.5;"><a contenteditable="false" href="#{entry['target_id']}" onclick="event.preventDefault(); const t = document.querySelector('[data-block-id=\\'{entry['target_id']}\\']'); if(t) t.scrollIntoView({{behavior: 'smooth', block: 'center'}});" style="color: #D4A373; text-decoration: none; font-weight: bold; cursor: pointer; border-bottom: 1px solid transparent; transition: all 0.2s; user-select: none;" onmouseover="this.style.borderBottom='1px solid #D4A373'; this.style.color='#3E2723';" onmouseout="this.style.borderBottom='1px solid transparent'; this.style.color='#D4A373';">{entry['text']}</a></span>"""
            
            toc_lines_html.append(link_html)

        combined_toc_content = "<br>".join(toc_lines_html)
        new_blocks.append(create_block("paragraph", combined_toc_content))

        new_blocks.append(create_block("paragraph", "<br>"))

        splice_index = req.target_line - 1
        if splice_index < 0: splice_index = 0
        if splice_index > len(blocks): splice_index = len(blocks)

        blocks[splice_index:splice_index] = new_blocks
        data["blocks"] = blocks
        data["metadata"]["last_modified"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)

        return {"status": "success"}

    except Exception as e:
        print(f"Error generating TOC: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    
class LayerNewDocRequest(BaseModel):
    base_filename: str
    suffix: str
    blocks: List[dict]

# Creates a new document by combining blocks from selected cognitive layers
@app.post("/api/docstral/layers/new_doc")
async def create_doc_from_layers(req: LayerNewDocRequest):
    """Creates a new document from selected layers"""
    base_name = req.base_filename.replace(".jdoc.json", "")
    
    import re
    base_name = re.sub(r'\s+', '_', base_name.strip())
    
    if req.suffix:
        new_filename = f"{base_name}_{req.suffix}.jdoc.json"
    else:
        new_filename = f"{base_name}.jdoc.json"
        
    counter = 1
    target_path = os.path.join(PLANCHETTE_DIR, new_filename)
    
    while os.path.exists(target_path):
        if req.suffix:
            new_filename = f"{base_name}_{req.suffix}_{counter}.jdoc.json"
        else:
            new_filename = f"{base_name}_{counter}.jdoc.json"
        target_path = os.path.join(PLANCHETTE_DIR, new_filename)

    doc_id = f"doc-{int(datetime.now().timestamp())}"
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    new_doc = {
        "file_version": "1",
        "metadata": {
            "id": doc_id,
            "name": new_filename.replace(".jdoc.json", "").replace("_", " "),
            "created_at": now_str,
            "last_modified": now_str
        },
        "blocks": req.blocks
    }

    try:
        with open(target_path, "w", encoding="utf-8") as f:
            json.dump(new_doc, f, indent=4, ensure_ascii=False)
        return {"status": "success", "filename": new_filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class VersionSaveRequest(BaseModel):
    filename: str
    version_name: str
    blocks: List[dict]
    visible_layers: List[str]

# Saves a new version of blocks and visible layers inside the current document's metadata
@app.post("/api/docstral/versions/save")
async def save_version(req: VersionSaveRequest):
    """Saves a new version inside the current document"""
    filepath = os.path.join(PLANCHETTE_DIR, req.filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)

        if "versions" not in data:
            data["versions"] = {}

        now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        data["versions"][req.version_name] = {
            "created_at": now_str,
            "last_modified": now_str,
            "visible_layers": req.visible_layers,
            "blocks": req.blocks
        }

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)

        return {"status": "success", "versions": list(data["versions"].keys())}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Returns a list of all stored versions for a given document file
@app.get("/api/docstral/versions/list")
async def list_versions(filename: str):
    """Returns a list of versions for a given file"""
    filepath = os.path.join(PLANCHETTE_DIR, filename)
    if not os.path.exists(filepath):
        return {"status": "success", "versions": []}

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
        versions = data.get("versions", {})
        return {"status": "success", "versions": list(versions.keys())}
    except Exception as e:
        return {"status": "success", "versions": []}

class VersionLoadRequest(BaseModel):
    filename: str
    version_name: str

# Loads and returns the specific blocks and visible layers for a designated document version
@app.post("/api/docstral/versions/load")
async def load_version(req: VersionLoadRequest):
    """Returns the blocks and visible layers for a specific version"""
    filepath = os.path.join(PLANCHETTE_DIR, req.filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        versions = data.get("versions", {})
        
        if req.version_name not in versions:
            raise ValueError(f"Version '{req.version_name}' not found inside file")

        v_data = versions[req.version_name]
        return {
            "status": "success", 
            "blocks": v_data.get("blocks", []),
            "visible_layers": v_data.get("visible_layers", [])
        }
    except ValueError as ve:
        raise HTTPException(status_code=404, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
class VersionDeleteRequest(BaseModel):
    filename: str
    version_name: str

# Deletes a specified version entirely from the document's version history
@app.post("/api/docstral/versions/delete")
async def delete_version(req: VersionDeleteRequest):
    """Deletes a specific version of the file"""
    filepath = os.path.join(PLANCHETTE_DIR, req.filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found")

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)

        if "versions" in data and req.version_name in data["versions"]:
            del data["versions"][req.version_name]
            
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
                
            return {"status": "success"}
        else:
            return {"status": "error", "message": "Version not found"}
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Saves the provided notes content into the designated scriptoria notes JSON file
@app.post("/api/notes/save")
async def save_notes(req: dict):
    os.makedirs(SCRIPTORIA_DIR, exist_ok=True)
    
    content_text = req.get("content", "")
    clean_notes = {"content": content_text}
    
    with open(SCRIPTORIA_NOTES, "w", encoding="utf-8") as f:
        json.dump(clean_notes, f, indent=4, ensure_ascii=False)
    return {"status": "success"}

# Loads and returns the content from the scriptoria notes JSON file
@app.get("/api/notes/load")
async def load_notes():
    if os.path.exists(SCRIPTORIA_NOTES):
        try:
            with open(SCRIPTORIA_NOTES, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and "content" in data:
                    return data
                return {"content": ""}
        except Exception:
            return {"content": ""}
    return {"content": ""}

# Parses and saves the current state of dataset rows into the JSONL dataset file
@app.post("/api/dataset/save")
async def save_dataset(req: dict):
    os.makedirs(SCRIPTORIA_DIR, exist_ok=True)
    with open(SCRIPTORIA_DATASET, "w", encoding="utf-8") as f:
        for row in req.get("rows", []):
            clean_row = {
                "instruction": row.get("instruction", ""),
                "input": row.get("input", ""),
                "output": row.get("output", "")
            }
            f.write(json.dumps(clean_row, ensure_ascii=False) + "\n")
    return {"status": "success"}

# Reads and returns the collection of rows from the scriptoria dataset JSONL file
@app.get("/api/dataset/load")
async def load_dataset():
    rows = []
    if os.path.exists(SCRIPTORIA_DATASET):
        with open(SCRIPTORIA_DATASET, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    try:
                        row_data = json.loads(line)
                        rows.append({
                            "id": row_data.get("id", f"row_{uuid.uuid4().hex[:12]}"),
                            "instruction": row_data.get("instruction", ""),
                            "input": row_data.get("input", ""),
                            "output": row_data.get("output", "")
                        })
                    except Exception: 
                        continue
    return {"rows": rows}

# Handles the uploading of multiple DOCX files specifically for scriptoria processing
@app.post("/api/docx/upload")
async def upload_multiple_files(files: List[UploadFile] = File(...)):
    saved_files = []
    for file in files:
        if not file.filename:
            continue
        safe_filename = os.path.basename(file.filename)
        ext = safe_filename.split('.')[-1].lower()
        
        if ext == 'jsonl':
            path = SCRIPTORIA_DATASET
        else:
            path = os.path.join(SCRIPTORIA_DOCX, safe_filename)
            
        with open(path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        saved_files.append(safe_filename)
        
    return {"status": "success", "filenames": saved_files}

# Removes a specific DOCX file from the scriptoria converters directory
@app.delete("/api/docx/delete/{filename}")
async def delete_docx(filename: str):
    safe_filename = os.path.basename(filename)
    path = os.path.join(SCRIPTORIA_DOCX, safe_filename)
    
    if os.path.exists(path):
        os.remove(path)
        return {"status": "success"}
    raise HTTPException(status_code=404, detail="File not found")

# Deletes a specific JSONL converted file from the system
@app.delete("/api/jsonl/delete/{filename}")
async def delete_jsonl_file(filename: str):
    safe_filename = os.path.basename(filename)
    deleted = False

    target_dirs = [
        JSONL_FILES_DIR,
        JSONL_MERGED_DIR
    ]

    for target_dir in target_dirs:
        for root, dirs, files in os.walk(target_dir):
            if safe_filename in files:
                file_path = os.path.join(root, safe_filename)
                try:
                    os.remove(file_path)
                    deleted = True
                    print(f"Deleted file: {file_path}")
                except Exception as e:
                    raise HTTPException(status_code=500, detail=f"Could not delete file: {str(e)}")

    if not deleted:
        raise HTTPException(status_code=404, detail="File not found on disk.")

    return {"status": "success", "message": f"File {safe_filename} deleted successfully."}

# Executes a specific Python script converter found in the converters directory
@app.post("/api/convert/{script_name}")
async def run_converter(script_name: str):
    script_path = os.path.join(CONVERTERS_DIR, f"{script_name}.py")
    if not os.path.exists(script_path):
        raise HTTPException(status_code=404, detail=f"Script {script_name}.py not found.")
    
    try:
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        
        result = subprocess.run(
            [sys.executable, script_path], 
            capture_output=True,
            env=env
        )
        
        stdout_text = result.stdout.decode('utf-8', errors='replace')
        stderr_text = result.stderr.decode('utf-8', errors='replace')

        if result.returncode != 0:
            print("--- BACKEND PYTHON ERROR LOG ---")
            print(stderr_text)
            raise HTTPException(status_code=500, detail=f"Script execution error: {stderr_text.strip()}")
            
        return {"status": "success", "message": "Conversion completed."}
        
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Generates a list of all successfully converted JSONL files across various directories
@app.get("/api/jsonl/list-all")
async def list_all_converted_files():
    jsonl_files = []
    
    search_patterns = [
        os.path.join(CONVERTERS_DIR, "jsonl_files", "**", "*.jsonl"),
        os.path.join(CONVERTERS_DIR, "jsonl_merged_files", "**", "*.jsonl")
    ]
    
    for pattern in search_patterns:
        for file_path in glob.glob(pattern, recursive=True):
            filename = os.path.basename(file_path)
            jsonl_files.append({"name": filename})
            
    unique_files = {f["name"]: f for f in jsonl_files}.values()
    
    return list(unique_files)

# Provides an endpoint to download a specific JSONL file based on its filename
@app.get("/api/jsonl/download/{filename}")
async def download_specific_jsonl(filename: str):
    safe_filename = os.path.basename(filename)
    found_path = None
    
    for root, dirs, files in os.walk(CONVERTERS_DIR):
        if safe_filename in files:
            found_path = os.path.join(root, safe_filename)
            break
            
    if not found_path or not os.path.exists(found_path):
        raise HTTPException(status_code=404, detail="The file was not physically found on the disk.")
        
    return FileResponse(found_path, media_type="application/jsonl", filename=safe_filename)

# Archives all converted JSONL files into a ZIP buffer and initiates a bulk download
@app.get("/api/jsonl/download-all")
async def download_all_converted():
    zip_buffer = BytesIO()
    has_files = False
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(CONVERTERS_DIR):
            for file in files:
                if file.endswith('.jsonl'):
                    has_files = True
                    file_path = os.path.join(root, file)
                    zf.write(file_path, arcname=file)
                    
    if not has_files:
        raise HTTPException(status_code=404, detail="No generated JSONL files found for archiving.")
        
    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer, 
        media_type="application/zip", 
        headers={"Content-Disposition": "attachment; filename=all_converted_files.zip"}
    )

# Completely removes the current scriptoria dataset file from the file system
@app.post("/api/dataset/clear")
async def clear_all_dataset():
    if os.path.exists(SCRIPTORIA_DATASET):
        os.remove(SCRIPTORIA_DATASET)
    return {"status": "success"}

# Imports a dataset file by handling its upload and processing
@app.post("/api/dataset/import-file")
async def import_dataset_file(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")
        
    filename = file.filename
    ext = filename.split('.')[-1].lower()
    
    if ext == 'jsonl':
        try:
            content = await file.read()
            text_content = content.decode('utf-8')
            
            rows = []
            for line in text_content.splitlines():
                if line.strip():
                    try:
                        row_data = json.loads(line)
                        rows.append({
                            "instruction": row_data.get("instruction", ""),
                            "input": row_data.get("input", ""),
                            "output": row_data.get("output", "")
                        })
                    except Exception:
                        continue
            
            with open(SCRIPTORIA_DATASET, "w", encoding="utf-8") as f:
                for row in rows:
                    f.write(json.dumps(row, ensure_ascii=False) + "\n")
                    
            return {"status": "success", "rows": rows}
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"JSONL quick import error: {str(e)}")

    elif ext == 'docx':
        try:
            content = await file.read()
            doc = docx.Document(BytesIO(content))
            
            dataset = []
            current_instruction = ""
            current_input = []
            current_output = []
            
            for para in doc.paragraphs:
                style_name = para.style.name.lower()
                text = para.text.strip()
                
                if not text:
                    continue
                    
                if "heading" in style_name or "title" in style_name:
                    if current_instruction or current_input or current_output:
                        dataset.append({
                            "instruction": current_instruction,
                            "input": "\n".join(current_input).strip(),
                            "output": "\n".join(current_output).strip()
                        })
                    current_instruction = text
                    current_input = []
                    current_output = []
                
                else:
                    if text.startswith("//"):
                        current_input.append(text[2:].strip())
                    elif text.startswith("*") or text.startswith("-"):
                        current_input.append(text[1:].strip())
                    else:
                        current_output.append(text)
                    
            if current_instruction or current_input or current_output:
                dataset.append({
                    "instruction": current_instruction,
                    "input": "\n".join(current_input).strip(),
                    "output": "\n".join(current_output).strip()
                })
            
            with open(SCRIPTORIA_DATASET, "w", encoding="utf-8") as f:
                for row in dataset:
                    f.write(json.dumps(row, ensure_ascii=False) + "\n")
                    
            return {"status": "success", "rows": dataset}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX parsing error: {str(e)}")
            
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format. Use .jsonl or .docx")

class GenerateDatasetRequest(BaseModel):
    prompt: str

# Instructs the AI model to generate and return a strict JSON array of dataset examples
@app.post("/api/dataset/generate-ai")
async def ai_generate_dataset(req: GenerateDatasetRequest):
    """
    Instructing the AI ​​model to return a strict JSON array with dataset examples.
    """

    system_instruction = """
    You are an expert dataset generator.
    The user will give you a topic or an instruction.
    You must generate diverse, high-quality examples.

    You MUST output ONLY a raw JSON array of objects in this exact structure,
    with no markdown formatting like ```json:

    [
        {
            "instruction": "The task or question",
            "input": "Any context or input string (leave empty if none)",
            "output": "The desired answer"
        }
    ]
    """

    messages = [
        ("system", system_instruction),
        ("human", req.prompt)
    ]

    try:
        llm = engine.get_llm("mistral", temperature=0.7)

        result = llm.invoke(messages)

        ai_text = (
            result.content
            if hasattr(result, "content")
            else str(result)
        )

        ai_text = ai_text.strip()

        ai_text = re.sub(
            r"^```json\s*",
            "",
            ai_text,
            flags=re.IGNORECASE
        )

        ai_text = re.sub(
            r"```$",
            "",
            ai_text
        )

        ai_text = ai_text.strip()

        new_rows = json.loads(ai_text)

        return {
            "status": "success",
            "new_rows": new_rows
        }

    except json.JSONDecodeError:
        return {
            "status": "error",
            "message": "AI model did not return valid JSON."
        }

    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }

# Exports the dataset into a requested format, validating the format type and dataset existence    
@app.get("/api/dataset/export/{format_type}")
async def export_dataset(format_type: str):
    if not os.path.exists(SCRIPTORIA_DATASET):
        raise HTTPException(status_code=404, detail="Dataset not found")

    rows = []
    with open(SCRIPTORIA_DATASET, "r", encoding="utf-8") as f:
        for line in f:
            if line.strip():
                rows.append(json.loads(line))

    exported_data = []

    if format_type == "4-4-2 IFT":
        def to_L(r): 
            return {"instruction": r.get('instruction', ''), "input": r.get('input', ''), "output": ""}
        
        def to_U(r): 
            return {"instruction": "", "input": r.get('input', ''), "output": r.get('output', '')}
        
        def to_W(r): 
            return {"instruction": r.get('instruction', ''), "input": r.get('input', ''), "output": r.get('output', '')}

        PATTERN = ['L', 'U', 'W', 'U', 'W', 'L', 'W', 'U', 'L']

        for step in PATTERN:
            for r in rows:
                if step == 'L':
                    exported_data.append(to_L(r))
                elif step == 'U':
                    exported_data.append(to_U(r))
                elif step == 'W':
                    exported_data.append(to_W(r))
    
    else:
        for r in rows:
            if format_type == "LMFT":
                exported_data.append({"prompt": f"{r.get('instruction', '')} {r.get('input', '')}".strip(), "completion": r.get('output', '')})
            elif format_type == "IFT":
                exported_data.append({"instruction": r.get('instruction', ''), "input": "", "output": r.get('output', '')})
            elif format_type == "4-L":
                exported_data.append({"instruction": r.get('instruction', ''), "input": r.get('input', ''), "output": ""})
            elif format_type == "4-U":
                exported_data.append({"instruction": "", "input": r.get('input', ''), "output": r.get('output', '')})
            elif format_type == "2-W":
                exported_data.append({"instruction": r.get('instruction', ''), "input": r.get('input', ''), "output": r.get('output', '')})

    output_content = "\n".join([json.dumps(obj, ensure_ascii=False) for obj in exported_data])
    
    return StreamingResponse(
        BytesIO(output_content.encode("utf-8")),
        media_type="application/jsonl",
        headers={"Content-Disposition": f"attachment; filename=export_{format_type}.jsonl"}
    )

class ClearLayerRequest(BaseModel):
    layer_key: str
    session_file: str

# Clears the session data from a specific cognitive layer when a frontend tab is closed
@app.post("/api/cognitive_layers/clear_session")
def clear_layer_session(req: ClearLayerRequest):
    engine.clear_session_from_layer(req.layer_key, req.session_file)
    return {"status": "success"}

# Calculates and returns statistics for Dharmachakra by reading the lifecycle meta-file
@app.get("/api/system_layers/stats")
def get_system_layer_stats():
    """Calculates the statistics for Dharmachakra, reading the cycle from the meta-file."""
    
    meta_path = os.path.join(engine.layers_dir, "system_meta.json")
    try:
        with open(meta_path, 'r', encoding='utf-8') as f:
            current_lifecycle_num = json.load(f).get("current_lifecycle", 1)
    except:
        current_lifecycle_num = 1

    passed_lifecycles = current_lifecycle_num - 1

    samsara_path = os.path.join(engine.layers_dir, engine.LAYER_MAP["samsara"], "samsara.json")
    try:
        with open(samsara_path, 'r', encoding='utf-8') as f:
            current_cycle_count = len([e for e in json.load(f) if "prompt" in e])
    except: 
        current_cycle_count = 0

    akasha_path = os.path.join(engine.layers_dir, engine.LAYER_MAP["akasha"], "akasha.json")
    try:
        with open(akasha_path, 'r', encoding='utf-8') as f: akasha_data = json.load(f)
    except: akasha_data = []

    real_entries = [e for e in akasha_data if "prompt" in e]
    total_prompts_all_time = len(real_entries)
    
    remaining_prompts = max(0, 1000 - current_cycle_count)
    progress_percentage = (current_cycle_count / 1000) * 100

    total_prompt_words = sum(len(e.get("prompt", "").split()) for e in real_entries)
    total_response_words = sum(len(e.get("response", "").split()) for e in real_entries)

    return {
        "lifecycle_number": current_lifecycle_num,
        "passed_lifecycles": passed_lifecycles, 
        "remaining_prompts": remaining_prompts,
        "progress_percentage": round(progress_percentage, 1),
        "total_prompts_all_time": total_prompts_all_time,
        "total_prompt_words": total_prompt_words,
        "total_response_words": total_response_words
    }

# Fully deletes Karma and Samsara layers and resets the lifecycle meta-file
@app.post("/api/system_layers/nirjara")
def execute_nirjara_reset():
    """Deletes Karma, Samsara layers and restarts the lifecycle meta-file."""
    try:
        for layer in ["samsara", "karma"]:
            folder_name = engine.LAYER_MAP[layer]
            json_path = os.path.join(engine.layers_dir, folder_name, f"{layer}.json")
            with open(json_path, 'w', encoding='utf-8') as f: json.dump([], f)
        
        maat_path = os.path.join(engine.layers_dir, engine.LAYER_MAP["samsara"], "hall_of_maat.json")
        with open(maat_path, 'w', encoding='utf-8') as f: json.dump([], f)

        meta_path = os.path.join(engine.layers_dir, "system_meta.json")
        with open(meta_path, 'w', encoding='utf-8') as f:
            json.dump({"current_lifecycle": 1}, f, indent=4)

        karma_vstore = engine.get_layer_vectorstore("karma")
        if karma_vstore:
            existing_data = karma_vstore.get()
            if existing_data and existing_data.get('ids'):
                karma_vstore.delete(ids=existing_data['ids'])
            
        print("[NIRJARA] Layers were reset. Lifecycle restarted at № 1.")
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Universal parser function designed to extract text content from various file formats    
def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Universal parser for extracting text from various formats."""
    ext = filename.split('.')[-1].lower()
    
    if ext in ['txt', 'json', 'jsonl', 'html', 'css', 'js', 'md', 'csv']:
        return file_bytes.decode('utf-8', errors='ignore')
        
    elif ext == 'docx':
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs])
        
    elif ext in ['xls', 'xlsx']:
        wb = openpyxl.load_workbook(filename=io.BytesIO(file_bytes), data_only=True)
        text_lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                line = " ".join([str(c) for c in row if c is not None])
                if line.strip(): text_lines.append(line)
        return "\n".join(text_lines)
        
    return file_bytes.decode('utf-8', errors='ignore')

# Creates a new custom cognitive layer as a JSON file specifically for In-Memory RAG
@app.post("/api/cognitive_layers/create_custom")
async def create_custom_layer(name: str = Form(...), files: List[UploadFile] = File(...)):
    """Creates a new custom layer as a JSON file (for In-Memory RAG)."""
    try:
        # Normalizes a string into an ASCII-only slug, substituting spaces with underscores
        import unicodedata
        def slugify(s):
            s = unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode('ascii')
            return "".join(c for c in s if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_").lower()

        clean_name = slugify(name)
        layer_key = f"custom_{clean_name}"
        
        layer_dir = os.path.join(engine.layers_dir, layer_key)
        os.makedirs(layer_dir, exist_ok=True)
        
        engine.LAYER_MAP[layer_key] = layer_key
        
        all_text = ""
        for file in files:
            content = await file.read()
            text = extract_text_from_file(content, file.filename)
            all_text += f"\n\n--- Source: {file.filename} ---\n{text}"
            
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        chunks = text_splitter.split_text(all_text)
        
        json_data = []
        for i, chunk in enumerate(chunks):
            json_data.append({
                "text": f"[CUSTOM LAYER: {name.upper()}]\n{chunk}",
                "metadata": {"source": name, "chunk": i}
            })
            
        json_path = os.path.join(layer_dir, f"{layer_key}.json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=4)
            
        return {"status": "success", "layer": layer_key}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Helper function to modify file permissions, allowing read-only files to be deleted
def remove_readonly(func, path, excinfo):
    """Helper function to remove read-only attributes to allow deletion."""
    os.chmod(path, stat.S_IWRITE)
    func(path)

# Instantly deletes a custom layer while handling URL decoding and file locks
@app.delete("/api/cognitive_layers/custom/{layer_key}")
def delete_custom_layer(layer_key: str):
    """Instantly delete a custom layer (without Windows File Locks)."""
    clean_key = urllib.parse.unquote(layer_key)
    if not clean_key.startswith("custom_"):
        clean_key = f"custom_{clean_key}"
        
    layer_dir = os.path.join(engine.layers_dir, clean_key)
    
    if clean_key in engine.LAYER_MAP:
        del engine.LAYER_MAP[clean_key]
    
    gc.collect()
    
    if os.path.exists(layer_dir):
        try:
            import shutil
            shutil.rmtree(layer_dir, onerror=remove_readonly)
            print(f"The custom layer {clean_key} was deleted successfully.")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Physical deletion error: {str(e)}")
            
    return {"status": "success"}
    
class SyncLayerRequest(BaseModel):
    layer_key: str       
    session_file: str     
    mode_id: int
    model_name: Optional[str] = None 

# Synchronizes an open tab's history directly into the Citta/Vritti layer, building its RAM vector
@app.post("/api/cognitive_layers/sync_session")
def sync_session_to_layer(req: SyncLayerRequest):
    """Syncs open tab history directly to the Citta/Vritti layer and forces Sleep if necessary."""
    try:
        engine.rebuild_dynamic_layer_vectors(req.mode_id, req.session_file)
        
        chosen_model = req.model_name
        if not chosen_model:
            try:
                models_data = get_ollama_models()
                if models_data and models_data.get("models"):
                    chosen_model = models_data["models"][0]
                else:
                    chosen_model = "mistral" 
            except:
                chosen_model = "mistral"

        layer_key = "citta" if req.mode_id in [3, 4] else "vritti"
        folder_name = engine.LAYER_MAP.get(layer_key)
        
        if folder_name:
            json_path = os.path.join(engine.layers_dir, folder_name, f"{layer_key}.json")
            try:
                with open(json_path, 'r', encoding='utf-8') as f:
                    dyn_data = json.load(f)
                
                session_raw = [item for item in dyn_data if item.get("session_file") == req.session_file and "distilled_sleep" not in item]
                
                if len(session_raw) >= 50:
                    print(f" [Magi] Background Sleep will use a model: '{chosen_model}'")
                    threading.Thread(
                        target=engine._perform_citta_sleep, 
                        args=(layer_key, req.session_file, chosen_model), 
                        daemon=True
                    ).start()
            except Exception as ex:
                print(f"The attempt to start sleep failed: {ex}")
                pass
                
        return {"status": "success", "message": f"Session synced to {req.layer_key.upper()}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class EndpointFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        return record.getMessage().find("/api/sessions/check-modified") == -1
logging.getLogger("uvicorn.access").addFilter(EndpointFilter())

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)